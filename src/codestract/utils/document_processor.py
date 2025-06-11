"""
Module for processing document files and extracting text content.
"""

import logging
import os
from typing import Optional

try:
    from docx import Document
    from openpyxl import load_workbook
    import fitz  # PyMuPDF
except ImportError as e:
    logging.warning(f"Document processing dependencies not available: {e}")
    Document = None
    load_workbook = None
    fitz = None


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Extract text content from a .docx file.
    
    Args:
        file_path: Path to the .docx file
        
    Returns:
        Extracted text content or None if extraction fails
    """
    if Document is None:
        logging.error("python-docx not available for .docx processing")
        return None
        
    try:
        doc = Document(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content)
        
    except Exception as e:
        logging.error(f"Error extracting text from {file_path}: {e}")
        return None


def extract_text_from_xlsx(file_path: str) -> Optional[str]:
    """
    Extract text content from a .xlsx file.
    
    Args:
        file_path: Path to the .xlsx file
        
    Returns:
        Extracted text content or None if extraction fails
    """
    if load_workbook is None:
        logging.error("openpyxl not available for .xlsx processing")
        return None
        
    try:
        workbook = load_workbook(file_path, data_only=True)
        text_content = []
        
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            text_content.append(f"Sheet: {sheet_name}")
            text_content.append("-" * (len(sheet_name) + 7))
            
            # Find the data range
            max_row = worksheet.max_row
            max_col = worksheet.max_column
            
            for row in range(1, max_row + 1):
                row_data = []
                has_data = False
                
                for col in range(1, max_col + 1):
                    cell_value = worksheet.cell(row=row, column=col).value
                    if cell_value is not None:
                        row_data.append(str(cell_value))
                        has_data = True
                    else:
                        row_data.append("")
                
                if has_data:
                    text_content.append(" | ".join(row_data))
            
            text_content.append("")  # Add spacing between sheets
        
        return "\n".join(text_content)
        
    except Exception as e:
        logging.error(f"Error extracting text from {file_path}: {e}")
        return None


def is_document_file(file_path: str) -> bool:
    """
    Check if a file is a supported document file.
    
    Args:
        file_path: Path to the file to check
        
    Returns:
        True if the file is a supported document format
    """
    DOCUMENT_EXTENSIONS = {".docx", ".xlsx", ".pdf"}
    ext = os.path.splitext(file_path)[1].lower()
    return ext in DOCUMENT_EXTENSIONS


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    Extract text content from a .pdf file.
    
    Args:
        file_path: Path to the .pdf file
        
    Returns:
        Extracted text content or None if extraction fails
    """
    if fitz is None:
        logging.error("PyMuPDF not available for .pdf processing")
        return None
        
    try:
        doc = fitz.open(file_path)
        text_content = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            if text.strip():
                text_content.append(f"Page {page_num + 1}:")
                text_content.append("-" * (len(f"Page {page_num + 1}:") + 1))
                text_content.append(text.strip())
                text_content.append("")  # Add spacing between pages
        
        doc.close()
        return "\n".join(text_content)
        
    except Exception as e:
        logging.error(f"Error extracting text from {file_path}: {e}")
        return None


def extract_document_content(file_path: str) -> Optional[str]:
    """
    Extract text content from a document file.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Extracted text content or None if extraction fails
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".xlsx":
        return extract_text_from_xlsx(file_path)
    elif ext == ".pdf":
        return extract_text_from_pdf(file_path)
    else:
        logging.warning(f"Unsupported document format: {ext}")
        return None 